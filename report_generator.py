#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
report_generator.py — Génération du rapport Word (.docx) pour HammerPy Insight.
Contient WordReportGenerator.
"""

import os
from datetime import datetime

from utils import FLOW_UNITS, VOLUME_UNITS

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordReportGenerator:
    """
    Génère un rapport d'ingénierie hydraulique professionnel au format Word (.docx)
    à partir des données du projet et des résultats d'analyse.
    """

    COULEUR_TITRE_HEX   = (0x1f, 0x53, 0x8d)
    COULEUR_ALERTE_HEX  = (0xc0, 0x00, 0x00)
    COULEUR_OK_HEX      = (0x1a, 0x7a, 0x1a)
    COULEUR_WARN_HEX    = (0xc0, 0x80, 0x00)
    COULEUR_NA_HEX      = (0x88, 0x88, 0x88)

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _rgb(self, rgb_tuple: tuple) -> RGBColor:
        return RGBColor(*rgb_tuple)

    def _setup_styles(self):
        for section in self.doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.0)

    def _add_title(self, text: str, level: int = 1, color_hex: tuple | None = None):
        p = self.doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = self._rgb(color_hex or self.COULEUR_TITRE_HEX)

    def _add_separator(self):
        p = self.doc.add_paragraph()
        p.add_run("─" * 80).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def _add_key_value(self, key: str, value: str, bold_value: bool = False):
        p = self.doc.add_paragraph()
        run_key = p.add_run(f"{key} : ")
        run_key.bold = True
        run_val = p.add_run(str(value))
        run_val.bold = bold_value

    def _add_alert(self, text: str, alert_type: str = "warning"):
        color_map = {
            "warning": self.COULEUR_ALERTE_HEX,
            "error":   self.COULEUR_ALERTE_HEX,
            "ok":      self.COULEUR_OK_HEX,
        }
        icon_map = {
            "warning": "⚠  ATTENTION : ",
            "error":   "✘  ERREUR : ",
            "ok":      "✔  CONFORME : ",
        }
        p = self.doc.add_paragraph()
        run = p.add_run(icon_map.get(alert_type, "• ") + text)
        run.font.color.rgb = self._rgb(color_map.get(alert_type, (0, 0, 0)))
        run.bold = (alert_type in ("warning", "error"))

    def _diag_icon(self, status: str) -> str:
        return {
            "OK":   "✔",
            "WARN": "⚠",
            "FAIL": "✘",
            "NA":   "—",
        }.get(status, "•")

    def _diag_color(self, status: str) -> tuple:
        return {
            "OK":   self.COULEUR_OK_HEX,
            "WARN": self.COULEUR_WARN_HEX,
            "FAIL": self.COULEUR_ALERTE_HEX,
            "NA":   self.COULEUR_NA_HEX,
        }.get(status, (0, 0, 0))

    def _insert_chart(self, png_path: str, width_cm: float = 15):
        """Insère une image PNG centrée dans le document Word."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(png_path, width=Cm(width_cm))
        except Exception as exc:
            err_p = self.doc.add_paragraph()
            err_run = err_p.add_run(f"[Graphique non insérable : {exc}]")
            err_run.italic = True
            err_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        self.doc.add_paragraph()

    def generate(self,
                 metadata: dict,
                 steady: dict | None,
                 transient: dict | None,
                 pn_label: str,
                 pn_value: float,
                 pmin_label: str,
                 pmin_value: float,
                 chart_png_path: str | None,
                 flow_unit: str = "m³/h",
                 volume_unit: str = "L",
                 volume_threshold_disp: float = 200.0,
                 workbook_summary: dict | None = None,
                 pump_summaries: list[dict] | None = None,
                 air_valve_data: dict | None = None,
                 diagnostics_checks: list[dict] | None = None,
                 diagnostics_summary: dict | None = None,
                 diag_kpi_chart_path: str | None = None,
                 diag_category_chart_path: str | None = None,
                 diag_compliance_chart_path: str | None = None,
                 diag_profile_chart_path: str | None = None) -> Document:
        """
        Génère le contenu complet du rapport Word.

        Args:
            metadata     : Informations du projet (nom_projet, ingenieur, date).
            steady       : Résultat du parser régime permanent.
            transient    : Résultat du parser analyse transitoire.
            pn_label     : Label de la classe PN (ex: "PN 16").
            pn_value     : Pression nominale en bar.
            pmin_label   : Label de la pression min admissible.
            pmin_value   : Valeur de la pression min en bar.
            flow_unit    : Unité d'affichage pour les débits ("m³/h" ou "L/s").
            volume_unit  : Unité d'affichage pour les volumes ("L" ou "m³").
            volume_threshold_disp : Seuil HPT affiché dans l'unité choisie.
            chart_png_path : Chemin vers l'image PNG du graphique.
            workbook_summary : Résumé du classeur HAMMER.
            pump_summaries   : Liste de résumés des pompes.
            air_valve_data   : Données ventouses/vidanges (profile, ventouses, vidanges).

        Returns:
            Le Document Word peuplé.
        """
        if steady and steady.get("success"):
            flow_raw = steady.get("flow_rate_m3h")
            if flow_raw is not None:
                flow_disp_val = flow_raw * FLOW_UNITS.get(flow_unit, 1.0)
                flow_disp_str = f"{flow_disp_val:.2f} {flow_unit}"
            else:
                flow_disp_str = "Non extrait – vérifier le fichier"
        else:
            flow_disp_str = "Non extrait – vérifier le fichier"

        if transient and transient.get("success"):
            vgas_raw = transient.get("max_gas_volume_l")
            if vgas_raw is not None:
                vgas_disp_val = vgas_raw / VOLUME_UNITS.get(volume_unit, 1.0)
                vgas_disp_str = f"{vgas_disp_val:.3f} {volume_unit}"
            else:
                vgas_disp_str = "—"
        else:
            vgas_raw = None
            vgas_disp_str = "—"

        threshold_disp_str = f"{volume_threshold_disp:.3f} {volume_unit}"
        now = metadata.get("date", datetime.now().strftime("%d/%m/%Y"))

        # ── En-tête du rapport ──────────────────────────────────────────
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run("NOTE TECHNIQUE D'ANALYSE HYDRAULIQUE\nCOUP DE BÉLIER — PROTECTION PAR RÉSERVOIR HPT")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self._rgb(self.COULEUR_TITRE_HEX)
        self.doc.add_paragraph()

        # ── Tableau de métadonnées du projet ───────────────────────────
        self._add_title("1. Identification du Projet", level=1)
        meta_table = self.doc.add_table(rows=4, cols=2)
        meta_table.style = "Table Grid"
        meta_data = [
            ("Nom du projet",        metadata.get("nom_projet", "—")),
            ("Ingénieur responsable",metadata.get("ingenieur",  "—")),
            ("Date d'établissement", now),
            ("Logiciel source",      "Bentley HAMMER – Résultats transitoires"),
        ]
        for i, (k, v) in enumerate(meta_data):
            row = meta_table.rows[i]
            cell_k = row.cells[0]
            cell_v = row.cells[1]
            cell_k.paragraphs[0].add_run(k).bold = True
            cell_v.paragraphs[0].add_run(v)

        self.doc.add_paragraph()

        # ── Régime Permanent ───────────────────────────────────────────
        self._add_separator()
        self._add_title("2. Caractérisation du Régime Permanent", level=1)

        if steady and steady.get("success"):
            flow = steady.get("flow_rate_m3h")
            hmt  = steady.get("hmt_m")
            self._add_key_value("Débit nominal de dimensionnement (Q)",
                                 flow_disp_str,
                                 bold_value=True)
            self._add_key_value("Hauteur Manométrique Totale (HMT)",
                                 f"{hmt} mCE" if hmt is not None else "Non extraite – vérifier le fichier",
                                 bold_value=True)
            self._add_key_value("Fichier station analysé",
                                 os.path.basename(metadata.get("station_filepath", "—")))
            self._add_key_value("Nombre de lignes lues", str(steady.get("n_rows", "—")))
        else:
            msg = steady.get("message", "—") if steady else "Fichier non chargé"
            self._add_alert(f"Régime permanent non disponible : {msg}", "error")

        self.doc.add_paragraph()

        # ── Modèle hydraulique (classeur HAMMER) ─────────────────────
        if workbook_summary:
            self._add_separator()
            self._add_title("2. Modèle Hydraulique (Classeur HAMMER)", level=1)

            p_intro = self.doc.add_paragraph(
                "Le modèle hydraulique a été extrait du classeur HAMMER (Flex Tables). "
                "Les données ci-dessous résument la structure du réseau analysé."
            )
            p_intro.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            self.doc.add_paragraph()

            counts = [
                ("Conduites (Pipes)", workbook_summary.get("pipes_count", 0)),
                ("Nœuds (Nodes)", workbook_summary.get("nodes_count", 0)),
                ("Pompes (Pumps)", workbook_summary.get("pumps_count", 0)),
                ("Réservoirs (Reservoirs)", workbook_summary.get("reservoirs_count", 0)),
                ("Réservoirs HPT", workbook_summary.get("hpt_count", 0)),
                ("Ventouses (Air Valves)", workbook_summary.get("air_valves_count", 0)),
            ]

            tbl = self.doc.add_table(rows=len(counts) + 1, cols=2)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdrs = ["Composant", "Nombre"]
            for j, h in enumerate(hdrs):
                run = tbl.rows[0].cells[j].paragraphs[0].add_run(h)
                run.bold = True
                run.font.color.rgb = self._rgb(self.COULEUR_TITRE_HEX)

            for i, (label, count) in enumerate(counts):
                tbl.rows[i + 1].cells[0].paragraphs[0].add_run(label)
                tbl.rows[i + 1].cells[1].paragraphs[0].add_run(str(count))

            self.doc.add_paragraph()

            materials = workbook_summary.get("materials", [])
            diam_min = workbook_summary.get("diameter_min_mm")
            diam_max = workbook_summary.get("diameter_max_mm")

            if materials or diam_min is not None:
                self._add_title("Propriétés des Conduites", level=2)
                if materials:
                    self._add_key_value("Matériaux détectés", ", ".join(materials))
                if diam_min is not None and diam_max is not None:
                    self._add_key_value("Diamètres",
                                        f"{diam_min} mm — {diam_max} mm")
                self.doc.add_paragraph()

            pmax_model = workbook_summary.get("pmax_bar")
            pmin_model = workbook_summary.get("pmin_bar")

            if pmax_model is not None or pmin_model is not None:
                self._add_title("Pressions Transitoires (Modèle)", level=2)
                if pmax_model is not None:
                    pmax_ok = pmax_model <= pn_value
                    self._add_key_value("P max transitoire (modèle)",
                                        f"{pmax_model:.2f} bar")
                    self._add_alert(
                        f"P max ({pmax_model:.2f} bar) ≤ {pn_label} ({pn_value} bar)"
                        if pmax_ok else
                        f"P max ({pmax_model:.2f} bar) > {pn_label} ({pn_value} bar) — DÉPASSEMENT",
                        "ok" if pmax_ok else "warning"
                    )
                if pmin_model is not None:
                    pmin_ok = pmin_model >= pmin_value
                    self._add_key_value("P min transitoire (modèle)",
                                        f"{pmin_model:.2f} bar")
                    self._add_alert(
                        f"P min ({pmin_model:.2f} bar) ≥ limite ({pmin_value} bar)"
                        if pmin_ok else
                        f"P min ({pmin_model:.2f} bar) < limite ({pmin_value} bar) — DÉPRESSION CRITIQUE",
                        "ok" if pmin_ok else "warning"
                    )
                self.doc.add_paragraph()

            vmax_pump = workbook_summary.get("vmax_pump_ls")
            if vmax_pump is not None:
                self._add_key_value("Débit max pompe (modèle)",
                                    f"{vmax_pump:.1f} L/s")

            self.doc.add_paragraph()

        # ── Données Pompe (multi-pompe) ────────────────────────────────
        if pump_summaries:
            self._add_separator()
            self._add_title("2b. Batterie de Pompes (Rapports Détaillés)", level=1)

            p_intro = self.doc.add_paragraph(
                "Les données ci-dessous proviennent des rapports détaillés pompe exportés depuis "
                "Bentley HAMMER. Elles caractérisent les points de fonctionnement nominaux et les "
                "courbes H(Q) des pompes analysées."
            )
            p_intro.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            self.doc.add_paragraph()

            for idx, ps in enumerate(pump_summaries, 1):
                if not ps or not ps.get("label") or ps["label"] == "—":
                    continue

                self._add_title(f"Pompe {idx}/{len(pump_summaries)} — {ps.get('label', '—')}", level=2)

                pump_rows = [
                    ("Identifiant",     str(ps.get("pump_id", "—"))),
                    ("Label",           str(ps.get("label", "—"))),
                    ("Conduite aval",   str(ps.get("downstream_pipe", "—"))),
                    ("Débit nominal (Q)",
                     f"{ps['flow_lps']:.1f} L/s" if ps.get("flow_lps") is not None else "—"),
                    ("HMT pompe",
                     f"{ps['pump_head_m']:.1f} m" if ps.get("pump_head_m") is not None else "—"),
                    ("Pression aspiration",
                     f"{ps['pressure_suction_bar']:.2f} bar" if ps.get("pressure_suction_bar") is not None else "—"),
                    ("Pression refoulement",
                     f"{ps['pressure_discharge_bar']:.2f} bar" if ps.get("pressure_discharge_bar") is not None else "—"),
                    ("NPSH disponible",
                     f"{ps['npsh_available_m']:.1f} m" if ps.get("npsh_available_m") is not None else "—"),
                    ("NPSH requis",
                     f"{ps['npsh_required_m']:.1f} m" if ps.get("npsh_required_m") is not None else "N/D"),
                    ("Points courbe H(Q)",
                     str(ps.get("n_curve_points", 0))),
                ]

                tbl = self.doc.add_table(rows=len(pump_rows) + 1, cols=2)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                hdrs = ["Paramètre", "Valeur"]
                for j, h in enumerate(hdrs):
                    run = tbl.rows[0].cells[j].paragraphs[0].add_run(h)
                    run.bold = True
                    run.font.color.rgb = self._rgb(self.COULEUR_TITRE_HEX)

                for i, (label, val) in enumerate(pump_rows):
                    tbl.rows[i + 1].cells[0].paragraphs[0].add_run(label)
                    tbl.rows[i + 1].cells[1].paragraphs[0].add_run(val)

                self.doc.add_paragraph()

                npsh_a = ps.get("npsh_available_m")
                if npsh_a is not None:
                    if npsh_a < 3:
                        self._add_alert(
                            f"NPSH disponible faible ({npsh_a:.1f} m) — Pompe {ps.get('label', '—')}. "
                            "Risque de cavitation. Vérifier l'altitude d'aspiration.",
                            "warning"
                        )
                    else:
                        self._add_alert(
                            f"NPSH disponible ({npsh_a:.1f} m) suffisant — Pompe {ps.get('label', '—')}.",
                            "ok"
                        )

                self.doc.add_paragraph()

        # ── Analyse Transitoire ────────────────────────────────────────
        self._add_separator()
        self._add_title("3. Synthèse de l'Analyse Transitoire (Coup de Bélier)", level=1)

        p_intro = self.doc.add_paragraph(
            "L'analyse transitoire par la méthode des caractéristiques a été réalisée via Bentley HAMMER. "
            "Les courbes enveloppes de pression maximale et minimale, ainsi que le volume d'air "
            "maximal dans le réservoir HPT, sont synthétisés ci-dessous."
        )
        p_intro.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        self.doc.add_paragraph()

        if transient and transient.get("success"):
            pmax = transient.get("max_pressure_bar")
            pmin = transient.get("min_pressure_bar")
            vgas = transient.get("max_gas_volume_l")
            cols_found = transient.get("critical_columns_found", [])
            is_sim     = transient.get("is_simulated", False)

            res_table = self.doc.add_table(rows=4, cols=3)
            res_table.style = "Table Grid"
            res_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdrs = ["Grandeur", "Valeur calculée", "Statut par rapport aux limites"]
            for j, h in enumerate(hdrs):
                cell = res_table.rows[0].cells[j]
                run  = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.color.rgb = self._rgb(self.COULEUR_TITRE_HEX)

            pmax_ok = (pmax is not None and pmax <= pn_value)
            res_table.rows[1].cells[0].paragraphs[0].add_run("Pression max. (Surpression)").bold = True
            res_table.rows[1].cells[1].paragraphs[0].add_run(f"{pmax} bar" if pmax is not None else "—")
            pmax_status_run = res_table.rows[1].cells[2].paragraphs[0].add_run(
                f"✔ OK (< {pn_label} = {pn_value} bar)" if pmax_ok
                else f"✘ DÉPASSEMENT – {pmax} bar > {pn_label} ({pn_value} bar)"
            )
            pmax_status_run.font.color.rgb = self._rgb(self.COULEUR_OK_HEX if pmax_ok else self.COULEUR_ALERTE_HEX)
            pmax_status_run.bold = not pmax_ok

            pmin_ok = (pmin is not None and pmin >= pmin_value)
            res_table.rows[2].cells[0].paragraphs[0].add_run("Pression min. (Dépression)").bold = True
            res_table.rows[2].cells[1].paragraphs[0].add_run(f"{pmin} bar" if pmin is not None else "—")
            pmin_label_txt = f"✔ OK (> limite = {pmin_value} bar)" if pmin_ok else f"✘ DÉPRESSION CRITIQUE – {pmin} bar < limite ({pmin_value} bar)"
            pmin_status_run = res_table.rows[2].cells[2].paragraphs[0].add_run(pmin_label_txt)
            pmin_status_run.font.color.rgb = self._rgb(self.COULEUR_OK_HEX if pmin_ok else self.COULEUR_ALERTE_HEX)
            pmin_status_run.bold = not pmin_ok

            vgas_ok = (vgas is not None and vgas <= volume_threshold_disp * VOLUME_UNITS.get(volume_unit, 1.0))
            res_table.rows[3].cells[0].paragraphs[0].add_run("Volume de gaz max. (HPT)").bold = True
            res_table.rows[3].cells[1].paragraphs[0].add_run(vgas_disp_str)
            vgas_status_run = res_table.rows[3].cells[2].paragraphs[0].add_run(
                f"✔ OK (≤ {threshold_disp_str})" if vgas_ok
                else f"✘ VOLUME ÉLEVÉ – {vgas_disp_str} > {threshold_disp_str}"
            )
            vgas_status_run.font.color.rgb = self._rgb(self.COULEUR_OK_HEX if vgas_ok else self.COULEUR_ALERTE_HEX)
            vgas_status_run.bold = not vgas_ok

            self.doc.add_paragraph()

            if is_sim:
                self._add_alert(
                    "Les colonnes critiques n'ont pas été trouvées dans le fichier. "
                    "Les valeurs affichées sont des PLACEHOLDERS de démonstration. "
                    "Vérifiez le format de votre export Bentley HAMMER.",
                    "warning"
                )
            else:
                self._add_alert(
                    f"Colonnes HAMMER identifiées : {', '.join(cols_found)}.",
                    "ok"
                )

        else:
            msg = transient.get("message", "—") if transient else "Fichier non chargé"
            self._add_alert(f"Données transitoires non disponibles : {msg}", "error")

        self.doc.add_paragraph()

        # ── Interprétation et recommandations ─────────────────────────
        self._add_separator()
        self._add_title("4. Interprétation et Diagnostic de Sécurité", level=1)

        if transient and transient.get("success"):
            pmax = transient.get("max_pressure_bar")
            pmin = transient.get("min_pressure_bar")
            vgas = transient.get("max_gas_volume_l")

            self._add_title("4.1  Surpression", level=2)
            if pmax is not None:
                if pmax > pn_value:
                    self._add_alert(
                        f"La surpression maximale ({pmax} bar) dépasse la pression nominale de la conduite "
                        f"({pn_label} = {pn_value} bar). Risque d'éclatement ou de détérioration accélérée des joints.\n"
                        "→ Recommandations : Augmenter la classe PN de la conduite ; réduire le débit de fermeture "
                        "des vannes (temps de fermeture plus long) ; installer une soupape de décharge.",
                        "warning"
                    )
                else:
                    self._add_alert(
                        f"Surpression maximale ({pmax} bar) inférieure à la PN choisie ({pn_label}). Situation sécurisée.",
                        "ok"
                    )

            self._add_title("4.2  Dépression", level=2)
            if pmin is not None:
                if pmin < pmin_value:
                    self._add_alert(
                        f"Une dépression critique ({pmin} bar) est détectée dans la conduite, en-dessous de la "
                        f"limite de sécurité définie ({pmin_value} bar).\n"
                        "Risques associés : intrusion d'air ou d'eau contaminée, collapsus de conduite "
                        "(PEHD/PVC), désamorçage des pompes.\n"
                        "→ Recommandations : Vérifier le pré-gonflage et le volume nominal du réservoir HPT ; "
                        "installer des ventouses automatiques aux points hauts ; ajuster le débit de démarrage des pompes.",
                        "warning"
                    )
                else:
                    self._add_alert(
                        f"Dépression minimale ({pmin} bar) dans les limites admissibles (>{pmin_value} bar).",
                        "ok"
                    )

            self._add_title("4.3  Réservoir HPT (Volume de Gaz)", level=2)
            if vgas is not None:
                threshold_l = volume_threshold_disp * VOLUME_UNITS.get(volume_unit, 1.0)
                if vgas > threshold_l:
                    self._add_alert(
                        f"Le volume d'air maximal dans le réservoir HPT ({vgas_disp_str}) est très élevé. "
                        "Cela peut indiquer un réservoir sous-dimensionné qui se vide totalement lors du transitoire, "
                        "perdant ainsi son efficacité.\n"
                        "→ Recommandation : Augmenter le volume nominal du réservoir HPT.",
                        "warning"
                    )
                else:
                    self._add_alert(
                        f"Volume de gaz maximal au HPT ({vgas_disp_str}) dans les limites admissibles (≤ {threshold_disp_str}).",
                        "ok"
                    )
        else:
            self.doc.add_paragraph("Aucune analyse disponible – charger un fichier transitoire valide.")

        self.doc.add_paragraph()

        # ── Graphique intégré ──────────────────────────────────────────
        if chart_png_path and os.path.exists(chart_png_path):
            self._add_separator()
            self._add_title("5. Visualisation Graphique — Courbes Enveloppes", level=1)
            self.doc.add_picture(chart_png_path, width=Cm(16))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        # ── Profil en Long — Ventaises & Vidanges (Phase 3) ──────────
        if air_valve_data:
            vents = air_valve_data.get("ventouses", [])
            drains = air_valve_data.get("vidanges", [])
            profile = air_valve_data.get("profile", [])
            pipe_dn = air_valve_data.get("pipe_dn_mm", 250)

            if vents or drains:
                self._add_separator()
                self._add_title("4. Profil en Long — Ventouses & Vidanges", level=1)

                p_intro = self.doc.add_paragraph(
                    f"Le profil en long de la conduite (DN {pipe_dn:.0f} mm) a été analysé "
                    f"pour localiser et dimensionner les ventouses et vidanges."
                )
                p_intro.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                self.doc.add_paragraph()

                if vents:
                    self._add_title(f"Ventouses recommandées ({len(vents)})", level=2)

                    tbl = self.doc.add_table(rows=len(vents) + 1, cols=4)
                    tbl.style = "Table Grid"
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                    hdrs = ["PK (m)", "Côte (m)", "Type", "DN (mm)"]
                    for j, h in enumerate(hdrs):
                        run = tbl.rows[0].cells[j].paragraphs[0].add_run(h)
                        run.bold = True
                        run.font.color.rgb = self._rgb((0x2d, 0x6a, 0x4f))

                    for i, v in enumerate(vents):
                        tbl.rows[i + 1].cells[0].paragraphs[0].add_run(
                            f"{v['pk_m']:.1f}")
                        tbl.rows[i + 1].cells[1].paragraphs[0].add_run(
                            f"{v['z_m']:.2f}")
                        tbl.rows[i + 1].cells[2].paragraphs[0].add_run(
                            v.get("type", "—"))
                        tbl.rows[i + 1].cells[3].paragraphs[0].add_run(
                            str(v.get("dn_mm", "—")))

                    self.doc.add_paragraph()

                if drains:
                    self._add_title(f"Vidanges recommandées ({len(drains)})", level=2)

                    tbl = self.doc.add_table(rows=len(drains) + 1, cols=6)
                    tbl.style = "Table Grid"
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                    hdrs = ["PK (m)", "Côte (m)", "Type", "DN",
                            "Dist. G (m)", "Dist. D (m)"]
                    for j, h in enumerate(hdrs):
                        run = tbl.rows[0].cells[j].paragraphs[0].add_run(h)
                        run.bold = True
                        run.font.color.rgb = self._rgb((0xef, 0x47, 0x6f))

                    for i, d in enumerate(drains):
                        tbl.rows[i + 1].cells[0].paragraphs[0].add_run(
                            f"{d['pk_m']:.1f}")
                        tbl.rows[i + 1].cells[1].paragraphs[0].add_run(
                            f"{d['z_m']:.2f}")
                        tbl.rows[i + 1].cells[2].paragraphs[0].add_run(
                            d.get("type", "—"))
                        tbl.rows[i + 1].cells[3].paragraphs[0].add_run(
                            str(d.get("dn_mm", "—")))
                        tbl.rows[i + 1].cells[4].paragraphs[0].add_run(
                            f"{d.get('distance_to_left_m', 0):.1f}")
                        tbl.rows[i + 1].cells[5].paragraphs[0].add_run(
                            f"{d.get('distance_to_right_m', 0):.1f}")

                    self.doc.add_paragraph()

                if not vents and not drains:
                    self._add_alert(
                        "Aucune ventouse ou vidange recommandée pour ce profil.",
                        "warning"
                    )

        # ── Diagnostic Système (Phase 4) ───────────────────────────
        if diagnostics_checks:
            self._add_separator()
            self._add_title("6. Diagnostic Système (Vérifications croisées)", level=1)

            # 6.1 Synthèse
            self._add_title("6.1  Synthèse des vérifications", level=2)
            if diagnostics_summary:
                n_ok   = diagnostics_summary.get("OK", 0)
                n_warn = diagnostics_summary.get("WARN", 0)
                n_fail = diagnostics_summary.get("FAIL", 0)
                n_na   = diagnostics_summary.get("NA", 0)
                n_tot  = diagnostics_summary.get("total", n_ok + n_warn + n_fail + n_na)
                self._add_key_value("Total vérifications", str(n_tot))
                p_intro = self.doc.add_paragraph()
                p_intro.add_run(
                    f"✔  OK : {n_ok}    ⚠  WARN : {n_warn}    ✘  FAIL : {n_fail}    —  N/A : {n_na}"
                ).bold = True
                if n_fail == 0 and n_warn == 0:
                    self._add_alert(
                        "Toutes les vérifications applicables sont conformes.",
                        "ok"
                    )
                elif n_fail > 0:
                    self._add_alert(
                        f"{n_fail} vérification(s) en échec — action corrective requise.",
                        "error"
                    )
                else:
                    self._add_alert(
                        f"{n_warn} avertissement(s) détecté(s) — à examiner.",
                        "warning"
                    )
                self.doc.add_paragraph()

            # 6.2 Détail par catégorie
            self._add_title("6.2  Détail par catégorie", level=2)
            current_cat = None
            for c in diagnostics_checks:
                cat = c.get("category", "")
                if cat != current_cat:
                    self._add_title(cat, level=3)
                    current_cat = cat

                status = c.get("status", "NA")
                code   = c.get("code", "")
                name   = c.get("name", "")
                msg    = c.get("message", "")
                val    = c.get("value")
                thresh = c.get("threshold")

                p = self.doc.add_paragraph()
                run_icon = p.add_run(self._diag_icon(status) + "  ")
                run_icon.font.color.rgb = self._rgb(self._diag_color(status))
                run_icon.bold = True

                run_code = p.add_run(f"[{code}]  ")
                run_code.bold = True
                run_code.font.color.rgb = self._rgb(self.COULEUR_TITRE_HEX)

                p.add_run(name)
                p.add_run(" — ")
                run_msg = p.add_run(msg)
                run_msg.font.color.rgb = self._rgb(self._diag_color(status))

                if val is not None or thresh is not None:
                    sub = self.doc.add_paragraph()
                    sub.paragraph_format.left_indent = Cm(0.8)
                    if val is not None:
                        sub_run = sub.add_run(f"Valeur : {val}")
                        sub_run.font.size = Pt(9)
                        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    if thresh is not None:
                        sub_run2 = sub.add_run(f"   Seuil : {thresh}")
                        sub_run2.font.size = Pt(9)
                        sub_run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # 6.3 Visualisations graphiques (Phase 4)
            self._add_title("6.3  Visualisations graphiques", level=2)
            self.doc.add_paragraph()

            # 6.3.a Synthèse (donut)
            if diag_kpi_chart_path and os.path.isfile(diag_kpi_chart_path):
                self._add_title("6.3.a  Synthèse globale", level=3)
                self._insert_chart(diag_kpi_chart_path, width_cm=15)

            # 6.3.b Par catégorie (barres empilées)
            if diag_category_chart_path and os.path.isfile(diag_category_chart_path):
                self._add_title("6.3.b  Résultats par catégorie (A → E)", level=3)
                self._insert_chart(diag_category_chart_path, width_cm=16)

            # 6.3.c Conformité HPT (3 barres)
            if diag_compliance_chart_path and os.path.isfile(diag_compliance_chart_path):
                self._add_title("6.3.c  Conformité paramètres transitoires", level=3)
                self._insert_chart(diag_compliance_chart_path, width_cm=16)

            # 6.3.d Profil en long (si profil chargé)
            if diag_profile_chart_path and os.path.isfile(diag_profile_chart_path):
                self._add_title("6.3.d  Profil en long & protections", level=3)
                self._insert_chart(diag_profile_chart_path, width_cm=16)

            # Si aucun graphe disponible
            if not any([
                diag_kpi_chart_path and os.path.isfile(diag_kpi_chart_path),
                diag_category_chart_path and os.path.isfile(diag_category_chart_path),
                diag_compliance_chart_path and os.path.isfile(diag_compliance_chart_path),
                diag_profile_chart_path and os.path.isfile(diag_profile_chart_path),
            ]):
                p_na = self.doc.add_paragraph()
                p_na.add_run(
                    "Aucun graphique disponible — données insuffisantes pour générer les visualisations."
                ).italic = True

        # ── Pied de page ───────────────────────────────────────────────
        self._add_separator()
        footer_p = self.doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(
            f"Document généré automatiquement par HammerPy Insight v3.0 — {now}"
        )
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(9)

        return self.doc
