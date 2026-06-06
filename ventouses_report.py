#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
ventouses_report.py — Génération du rapport Word (.docx) pour les
ventouses et vidanges, avec profil en long en image.

Crée un document structuré contenant :
  1. En-tête (projet, ingénieur, date, DN conduite)
  2. Profil en long (image PNG)
  3. Tableau ventouses (PK, côte, type, DN, pente)
  4. Tableau vidanges (PK, côte, type, DN, distances)
  5. Statistiques globales
"""

import os
import tempfile
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Pt = Cm = RGBColor = Inches = WD_ALIGN_PARAGRAPH = None


# Couleurs (R, G, B) en tuples
COULEUR_TITRE = (0x1f, 0x53, 0x8d)
COULEUR_VENTOUSE = (0x06, 0xd6, 0xa0)        # Vert
COULEUR_VIDANGE = (0xef, 0x47, 0x6f)         # Rouge
COULEUR_INFO = (0x4c, 0xc9, 0xf0)            # Bleu clair
COULEUR_ALERTE = (0xc0, 0x00, 0x00)
COULEUR_GRIS = (0x66, 0x66, 0x66)


class VentousesReportGenerator:
    """
    Génère un rapport Word (.docx) dédié aux ventouses et vidanges
    à partir d'un AirValveSizing et de la figure Matplotlib du profil.
    """

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx n'est pas installé. "
                "Lancez : pip install python-docx"
            )
        self.doc = Document()
        self._setup_styles()

    def _rgb(self, rgb_tuple: tuple) -> RGBColor:
        return RGBColor(*rgb_tuple)

    def _setup_styles(self):
        for section in self.doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.0)

    def _add_title(self, text: str, level: int = 1, color: tuple | None = None):
        p = self.doc.add_heading(text, level=level)
        if p.runs:
            p.runs[0].font.color.rgb = self._rgb(color or COULEUR_TITRE)
        return p

    def _add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                       color: tuple | None = None, alignment=None):
        p = self.doc.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = self._rgb(color)
        return p

    def _add_separator(self):
        p = self.doc.add_paragraph()
        run = p.add_run("─" * 80)
        run.font.color.rgb = self._rgb(COULEUR_GRIS)
        run.font.size = Pt(8)

    def _add_kv_table(self, rows: list[tuple[str, str]]) -> None:
        """Ajoute un tableau clé/valeur (2 colonnes) sans bordure."""
        table = self.doc.add_table(rows=len(rows), cols=2)
        table.autofit = True
        for i, (key, value) in enumerate(rows):
            cell_key = table.rows[i].cells[0]
            cell_val = table.rows[i].cells[1]
            # Clé en gras
            p_key = cell_key.paragraphs[0]
            r_key = p_key.add_run(key)
            r_key.bold = True
            # Valeur
            cell_val.paragraphs[0].add_run(str(value))

    def _add_ventouses_table(self, ventouses: list[dict]) -> None:
        """Ajoute le tableau des ventouses (4 colonnes)."""
        if not ventouses:
            self._add_paragraph("Aucune ventouse recommandée.",
                                italic=True, color=COULEUR_GRIS)
            return

        # En-tête
        table = self.doc.add_table(rows=1 + len(ventouses), cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        headers = ["PK (m)", "Côte Z (m)", "Pente (%)", "Type", "DN (mm)"]
        for i, h in enumerate(headers):
            p = hdr[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = self._rgb(COULEUR_VENTOUSE)

        # Données
        for i, v in enumerate(ventouses, start=1):
            cells = table.rows[i].cells
            cells[0].paragraphs[0].add_run(f"{v['pk_m']:.1f}")
            cells[1].paragraphs[0].add_run(f"{v['z_m']:.2f}")
            cells[2].paragraphs[0].add_run(f"{v.get('pente_pct', 0):.2f}")
            cells[3].paragraphs[0].add_run(str(v.get("type", "")))
            cells[4].paragraphs[0].add_run(str(v.get("dn_mm", "")))

    def _add_vidanges_table(self, vidanges: list[dict]) -> None:
        """Ajoute le tableau des vidanges (6 colonnes)."""
        if not vidanges:
            self._add_paragraph("Aucune vidange recommandée.",
                                italic=True, color=COULEUR_GRIS)
            return

        table = self.doc.add_table(rows=1 + len(vidanges), cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        headers = ["PK (m)", "Côte Z (m)", "Type", "DN (mm)",
                   "Dist. V. G (m)", "Dist. V. D (m)"]
        for i, h in enumerate(headers):
            p = hdr[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = self._rgb(COULEUR_VIDANGE)

        for i, d in enumerate(vidanges, start=1):
            cells = table.rows[i].cells
            cells[0].paragraphs[0].add_run(f"{d['pk_m']:.1f}")
            cells[1].paragraphs[0].add_run(f"{d['z_m']:.2f}")
            cells[2].paragraphs[0].add_run(str(d.get("type", "")))
            cells[3].paragraphs[0].add_run(str(d.get("dn_mm", "")))
            cells[4].paragraphs[0].add_run(f"{d.get('distance_to_left_m', 0):.1f}")
            cells[5].paragraphs[0].add_run(f"{d.get('distance_to_right_m', 0):.1f}")

    def _compute_stats(self, air_valve_sizer, metadata: dict) -> dict:
        """Calcule les statistiques globales du dimensionnement."""
        profile = air_valve_sizer.profile
        ventouses = air_valve_sizer.ventouses
        vidanges = air_valve_sizer.vidanges
        return {
            "nb_points_profil": len(profile),
            "longueur_totale_m": (profile[-1]["pk_m"] - profile[0]["pk_m"])
                                 if profile else 0.0,
            "cote_min_m": min((p["z_m"] for p in profile), default=None),
            "cote_max_m": max((p["z_m"] for p in profile), default=None),
            "nb_ventouses": len(ventouses),
            "nb_vidanges": len(vidanges),
            "nb_ventouses_simples": sum(
                1 for v in ventouses
                if "simple" in v.get("type", "").lower()
            ),
            "nb_ventouses_combinees": sum(
                1 for v in ventouses
                if "combinée" in v.get("type", "").lower()
                or "combinee" in v.get("type", "").lower()
            ),
            "nb_ventouses_grande_orifice": sum(
                1 for v in ventouses
                if "grande orifice" in v.get("type", "").lower()
            ),
            "dn_ventouse_max_mm": max((v["dn_mm"] for v in ventouses), default=None),
            "dn_ventouse_min_mm": min((v["dn_mm"] for v in ventouses), default=None),
        }

    def generate(
        self,
        air_valve_sizer,
        metadata: dict | None = None,
        profile_chart_png_path: str | None = None,
    ) -> "Document":
        """
        Génère le rapport complet et retourne l'objet Document (non sauvegardé).

        Args:
            air_valve_sizer : Instance de AirValveSizing
            metadata        : {nom_projet, ingenieur, date, dn_mm, ...}
            profile_chart_png_path : Chemin du PNG du graphique profil en long
                                     (si None, le graphique n'est pas inclus)

        Returns:
            Objet docx.Document (à sauvegarder via .save(filepath))
        """
        metadata = metadata or {}
        profile = air_valve_sizer.profile
        ventouses = air_valve_sizer.ventouses
        vidanges = air_valve_sizer.vidanges
        stats = self._compute_stats(air_valve_sizer, metadata)

        # ── 1. En-tête ─────────────────────────────────────────────
        nom_projet = metadata.get("nom_projet", "Projet sans nom")
        ingenieur = metadata.get("ingenieur", "—")
        date_rapport = metadata.get("date", datetime.now().strftime("%d/%m/%Y"))
        dn_mm = metadata.get("dn_mm", air_valve_sizer.pipe_dn_mm)

        self._add_title("Note Technique — Ventouses & Vidanges", level=0)
        self._add_separator()
        self._add_kv_table([
            ("Projet",           nom_projet),
            ("Ingénieur",        ingenieur),
            ("Date du rapport",  date_rapport),
            ("DN conduite",      f"{dn_mm} mm"),
            ("Source du profil", metadata.get("profil_source", "—")),
        ])

        # ── 2. Profil en long (image) ──────────────────────────────
        self._add_title("1. Profil en Long", level=1)
        if profile_chart_png_path and os.path.isfile(profile_chart_png_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(profile_chart_png_path, width=Cm(15.5))
        else:
            self._add_paragraph(
                "[Graphique du profil en long non disponible — "
                "veuillez charger un profil et générer le graphique]",
                italic=True, color=COULEUR_ALERTE,
            )

        # Légende
        self._add_paragraph(
            "Légende : marqueurs ▲ verts = ventouses recommandées, "
            "▼ rouges = vidanges recommandées.",
            italic=True, color=COULEUR_GRIS,
        )

        # ── 3. Caractéristiques du profil ─────────────────────────
        self._add_title("2. Caractéristiques du Profil", level=1)
        if profile:
            self._add_kv_table([
                ("Nombre de points",        f"{stats['nb_points_profil']}"),
                ("Longueur totale",         f"{stats['longueur_totale_m']:.1f} m"),
                ("Cote minimale (TN)",      f"{stats['cote_min_m']:.2f} m"
                                              if stats['cote_min_m'] is not None else "—"),
                ("Cote maximale (TN)",      f"{stats['cote_max_m']:.2f} m"
                                              if stats['cote_max_m'] is not None else "—"),
            ])
        else:
            self._add_paragraph("Aucun profil en long chargé.",
                                italic=True, color=COULEUR_ALERTE)

        # ── 4. Ventouses ───────────────────────────────────────────
        self._add_title(
            f"3. Ventouses Recommandées ({stats['nb_ventouses']})",
            level=1, color=COULEUR_VENTOUSE,
        )
        self._add_ventouses_table(ventouses)

        if ventouses:
            # Synthèse
            summary_rows = [
                ("Ventouses simples (anti-vide)",
                 f"{stats['nb_ventouses_simples']}"),
                ("Ventouses combinées (admission + dégazage)",
                 f"{stats['nb_ventouses_combinees']}"),
                ("Ventouses grande orifice (admission rapide)",
                 f"{stats['nb_ventouses_grande_orifice']}"),
                ("DN min recommandé",
                 f"{stats['dn_ventouse_min_mm']} mm"
                 if stats['dn_ventouse_min_mm'] is not None else "—"),
                ("DN max recommandé",
                 f"{stats['dn_ventouse_max_mm']} mm"
                 if stats['dn_ventouse_max_mm'] is not None else "—"),
            ]
            self._add_paragraph("")
            self._add_paragraph("Synthèse :", bold=True)
            self._add_kv_table(summary_rows)

        # ── 5. Vidanges ────────────────────────────────────────────
        self._add_title(
            f"4. Vidanges Recommandées ({stats['nb_vidanges']})",
            level=1, color=COULEUR_VIDANGE,
        )
        self._add_vidanges_table(vidanges)

        # ── 6. Notes méthodologiques ───────────────────────────────
        self._add_title("5. Méthodologie & Hypothèses", level=1)
        notes = [
            "Règles de dimensionnement appliquées (pré-dimensionnement) :",
            "  • Ventouse anti-vide (simple) : DN ≥ DN conduite / 12 (pentes > 3%)",
            "  • Ventouse combinée (admission + dégazage) : DN ≥ DN conduite / 10 (pentes < 0.5%)",
            "  • Ventouse grande orifice : DN ≥ DN conduite / 8 (pentes modérées 0.5%–3%)",
            "  • Vidange : DN ≥ DN conduite / 10, distance min 50 m aux ventouses, "
            "distance max 500 m entre vidanges.",
            "",
            "Hypothèses :",
            "  • Profil en long considéré en régime permanent (pas de transitoire).",
            "  • L'altitude est exprimée en mètres NGF (côte TN du terrain naturel).",
            "  • Les diamètres sont arrondis au DN standard supérieur le plus proche.",
        ]
        for line in notes:
            self._add_paragraph(line)

        # ── 7. Pied de page ────────────────────────────────────────
        self._add_separator()
        self._add_paragraph(
            f"Rapport généré automatiquement par HammerPy Insight — "
            f"{datetime.now().strftime('%d/%m/%Y à %H:%M')}",
            italic=True, color=COULEUR_GRIS,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        return self.doc

    def save(self, filepath: str) -> str:
        """
        Sauvegarde le document (doit avoir été généré via .generate()).

        Args:
            filepath: Chemin du .docx à créer
        Returns:
            Le même filepath (pour chaînage)
        """
        self.doc.save(filepath)
        return filepath


# ── Helper pour export direct depuis AirValveSizing ─────────────────

def export_ventouses_report(
    air_valve_sizer,
    output_path: str,
    metadata: dict | None = None,
    profile_chart_png_path: str | None = None,
) -> str:
    """
    Helper tout-en-un : génère et sauvegarde le rapport Ventouses/Vidanges.

    Args:
        air_valve_sizer : Instance de AirValveSizing
        output_path     : Chemin du .docx à créer
        metadata        : {nom_projet, ingenieur, date, dn_mm}
        profile_chart_png_path : PNG du profil en long (optionnel)

    Returns:
        Chemin du fichier créé
    """
    gen = VentousesReportGenerator()
    gen.generate(
        air_valve_sizer=air_valve_sizer,
        metadata=metadata or {},
        profile_chart_png_path=profile_chart_png_path,
    )
    gen.save(output_path)
    return output_path
